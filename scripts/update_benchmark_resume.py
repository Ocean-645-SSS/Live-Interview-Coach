from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SOURCE = Path("evaluation/materials/个人简历.docx")

PROJECT_PARAGRAPHS = [
    "Interview Coach｜可追溯中文技术模拟面试系统",
    "AI 应用开发                                      2026.01 - 至今",
    "简介：面向中文技术模拟面试的 AI 应用，围绕候选人资料、目标岗位、版本化题库与可选公司面试情报生成面试计划；支持实时语音面试、逐题结构化评价、可追溯报告与长期能力画像。",
    "技术栈：Python、FastAPI、LiveKit、LightRAG、PostgreSQL、Redis、SQLAlchemy、Alembic、MCP、Docker。",
    "·  基于 FastAPI 与 InterviewService 提供候选人画像、岗位资料、面试计划、Session、回答、评价、报告和准备任务 API；使用 PostgreSQL + SQLAlchemy + Alembic 持久化面试领域事实。",
    "·  通过独立 LightRAG 服务管理知识库生命周期、文档解析、索引和查询；以 kb_id 隔离元数据、原始文件、分块索引与查询审计，查询返回 context、references 与 ranked chunks。",
    "·  使用 LiveKit interview-agent 驱动实时音频、逐题流程控制、热词注入和回答提交；面试状态由版本化事件驱动，重复或过期事件被拒绝。",
    "·  基于题目 rubric 输出结构化评价，保存原始与规范化转写、逐题评价和 evidence refs；长期能力画像只从可追溯评价聚合训练建议。",
    "·  使用 Redis 与 Worker 承载准备类后台任务、队列协调和公司面试情报缓存；情报经过规范化、抽取与聚合后供规划器选择或个性化面试问题。",
]

SKILL_PARAGRAPHS = [
    "·   熟悉 Python，具备 Java 编程基础，掌握面向对象、集合、异常处理及多线程等知识。",
    "·   熟悉 FastAPI、LiveKit、PostgreSQL、Redis、Docker 等组件，能够参与本地部署的 AI 应用后端开发。",
    "·   熟悉以 LightRAG 为核心的知识库导入、文档解析、索引、结构化查询与证据审计链路，理解检索结果应以 context、references 和 ranked chunks 为依据。",
    "·   熟悉版本化题库、结构化 rubric 评价、面试状态机和可追溯能力画像的工程实现思路。",
]


def replace_skills(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "专业技能"
    )
    for paragraph, text in zip(
        paragraphs[start + 1 : start + 1 + len(SKILL_PARAGRAPHS)], SKILL_PARAGRAPHS, strict=True
    ):
        paragraph.text = text


def replace_project_experience(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "项目经历"
    )
    templates = [
        deepcopy(paragraph._p)
        for paragraph in paragraphs[start + 1 : start + 1 + len(PROJECT_PARAGRAPHS)]
    ]
    for paragraph in paragraphs[start + 1 :]:
        paragraph._element.getparent().remove(paragraph._element)

    anchor = document.paragraphs[start]._p
    for text, template in zip(PROJECT_PARAGRAPHS, templates, strict=True):
        anchor.addnext(template)
        paragraph = Paragraph(template, document._body)
        paragraph.text = text
        anchor = template


def main() -> None:
    document = Document(SOURCE)
    replace_skills(document)
    replace_project_experience(document)
    document.save(SOURCE)


if __name__ == "__main__":
    main()
